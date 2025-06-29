import asyncio
import base64
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

from docx import Document
from langchain_redis import RedisConfig, RedisVectorStore
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, Spacer
from scrapy.http import Request
from scrapy.linkextractors import LinkExtractor
from scrapy.spiders import Spider

from src.schemas.graph_schema import Model
from src.schemas.web_scraping_schema import OutputType
from src.services.web_scraper.scraper_config import (
    CONTENT_SELECTORS,
    EXCLUDED_SELECTORS,
    IGNORED_EXTENSIONS,
)
from src.utils.select_model import get_embedding_model


def install_reactor():
    """Install asyncio reactor if not already installed."""
    if "twisted.internet.reactor" not in sys.modules:
        import twisted.internet.asyncioreactor

        twisted.internet.asyncioreactor.install()


install_reactor()


class ScrapySpider(Spider):
    """
    Scrapy spider that respects domain restrictions and depth limits.
    """

    name = "spider"

    def __init__(
        self,
        start_urls=None,
        max_depth=1,
        allowed_domains=None,
        content_selectors=None,
        excluded_selectors=None,
        *args,
        **kwargs,
    ):
        super(ScrapySpider, self).__init__(*args, **kwargs)

        self.start_urls = start_urls or []
        self.max_depth = max_depth
        self.allowed_domains = allowed_domains or []

        self.content_selectors = content_selectors or CONTENT_SELECTORS
        self.excluded_selectors = excluded_selectors or EXCLUDED_SELECTORS

        self.scraped_data = {}
        self.failed_urls = set()

        self.link_extractor = LinkExtractor(
            allow_domains=self.allowed_domains if self.allowed_domains else None,
            deny_extensions=IGNORED_EXTENSIONS,
        )

    def parse(self, response):
        """
        Main parsing method that extracts content and follows links.
        """
        current_depth = response.meta.get("depth", 0)

        self._extract_content(response)

        if current_depth < self.max_depth:
            links = self.link_extractor.extract_links(response)

            for link in links:
                if self.allowed_domains:
                    link_domain = urlparse(link.url).netloc
                    if not any(
                        domain in link_domain for domain in self.allowed_domains
                    ):
                        continue

                yield Request(
                    url=link.url,
                    callback=self.parse,
                    meta={"depth": current_depth + 1},
                    errback=self.handle_error,
                    dont_filter=False,
                )

    def _extract_content(self, response):
        """
        Extract text content from the response and store it.
        First get content from CONTENT_SELECTORS, then remove EXCLUDED_SELECTORS.
        """
        try:
            title = response.css("title::text").get()
            title = title.strip() if title else "No title"

            content_parts = self._get_content_with_exclusions(response)

            full_content = f"Title: {title}\n\n"
            full_content += "Content:\n"
            full_content += "\n".join(content_parts[:200])

            self.scraped_data[response.url] = full_content
            self.logger.info(f"Successfully scraped content from: {response.url}")

        except Exception as e:
            self.logger.error(f"Error extracting content from {response.url}: {str(e)}")
            self.failed_urls.add(response.url)

    def _get_content_with_exclusions(self, response):
        """Helper method to extract content while excluding specified selectors."""
        content_parts = []

        self._extract_from_content_selectors(response, content_parts)

        return content_parts

    def _extract_from_content_selectors(self, response, content_parts):
        """Extract content from the main content selectors."""
        for selector in self.content_selectors:
            if not response.css(selector):
                continue

            filtered_text = self._filter_excluded_text(
                response, response.css(f"{selector} ::text").getall()
            )

            if filtered_text:
                content_parts.extend(filtered_text)

    def _filter_excluded_text(self, response, text_list):
        """Filter out text that appears in excluded selectors."""
        if not text_list:
            return []

        excluded_text = set()
        for excluded_selector in self.excluded_selectors:
            excluded_text.update(
                text.strip()
                for text in response.css(f"{excluded_selector} ::text").getall()
                if text.strip()
            )

        return [
            text.strip()
            for text in text_list
            if text.strip() and text.strip() not in excluded_text
        ]

    def handle_error(self, failure):
        """
        Handle request failures.
        """
        self.logger.error(
            f"Request failed: {failure.request.url} - {str(failure.value)}"
        )
        self.failed_urls.add(failure.request.url)

    def _save_scraped_content_to_file(self, output_type: str, output_path: str) -> str:
        """Save scraped content to file based on output type."""
        if not self.scraped_data:
            return ""

        output_dir = Path(output_path)
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        all_content = self._get_all_content()

        if output_type.lower() in [OutputType.TXT.value, "text"]:
            filename = self._save_as_text(str(output_dir), timestamp, all_content)
        elif output_type.lower() == OutputType.HTML.value:
            filename = self._save_as_html(str(output_dir), timestamp)
        elif output_type.lower() == OutputType.DOCX.value:
            filename = self._save_as_docx(str(output_dir), timestamp)
        elif output_type.lower() == OutputType.PDF.value:
            filename = self._save_as_pdf(str(output_dir), timestamp)
        elif output_type.lower() == OutputType.JSON.value:
            filename = self._save_as_json(str(output_dir), timestamp)
        elif output_type.lower() == OutputType.STRING.value:
            return all_content
        else:
            raise ValueError(f"Unsupported output type: {output_type}")

        return filename

    def _get_all_content(self) -> str:
        all_content = ""
        for url, content in self.scraped_data.items():
            all_content += f"URL: {url}\n\n{content}\n\n{'-'*80}\n\n"
        return all_content

    def _save_as_text(self, output_path, timestamp, all_content):
        filename = f"{output_path}/scraped_content_{timestamp}.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(all_content)
        return filename

    def _save_as_html(self, output_path, timestamp):
        filename = f"{output_path}/scraped_content_{timestamp}.html"
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Scraped Content</title>
</head>
<body>
    <h1>Scraped Content - {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</h1>
"""
        for url, content in self.scraped_data.items():
            html_content += f"""
    <div style="margin-bottom: 30px; border-bottom: 1px solid #ccc; padding-bottom: 20px;">
        <h2><a href="{url}" target="_blank">{url}</a></h2>
        <pre style="white-space: pre-wrap; font-family: Arial, sans-serif;">{content}</pre>
    </div>
"""
        html_content += "</body></html>"

        with open(filename, "w", encoding="utf-8") as f:
            f.write(html_content)
        return filename

    def _save_as_docx(self, output_path, timestamp):
        """Save scraped content as DOCX file."""
        filename = f"{output_path}/scraped_content_{timestamp}.docx"

        doc = Document()

        doc.add_heading("Scraped Content", 0)
        doc.add_paragraph(
            f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        doc.add_paragraph("")

        for i, (url, content) in enumerate(self.scraped_data.items()):
            if i > 0:
                doc.add_page_break()

            doc.add_heading(f"URL: {url}", level=1)
            doc.add_paragraph("")

            content_lines = content.split("\n")
            for line in content_lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")

        doc.save(filename)
        self.logger.debug(f"Successfully saved DOCX file: {filename}")
        return filename

    def _try_register_local_font(self, font_name, font_path):
        """Try to register a local TTF font."""
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            self.logger.debug(f"Using font: {font_name}")
            return font_name
        except Exception as e:
            self.logger.debug(f"Failed to register font {font_name}: {str(e)}")
            return None

    def _try_register_cid_font(self, cid_font_name):
        """Try to register a CID font."""
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        try:
            pdfmetrics.registerFont(UnicodeCIDFont(cid_font_name))
            self.logger.debug(f"Using built-in CID font {cid_font_name}")
            return cid_font_name
        except Exception:
            return None

    def _find_preferred_local_font(self, font_dir):
        """Find and register preferred fonts for Hungarian text."""
        preferred_fonts = ["DejaVuSans", "FreeSans", "NotoSans-Regular"]

        for preferred_name in preferred_fonts:
            font_path = font_dir / f"{preferred_name}.ttf"
            if font_path.exists():
                result = self._try_register_local_font(preferred_name, font_path)
                if result:
                    return result
        return None

    def _register_unicode_font(self):
        """Register a Unicode-capable font for PDF generation."""

        font_dir = Path(__file__).parent / "fonts"
        if not font_dir.exists():
            try:
                font_dir.mkdir(parents=True, exist_ok=True)
                self.logger.debug(f"Created fonts directory: {font_dir}")
            except Exception as e:
                self.logger.error(f"Failed to create fonts directory: {str(e)}")

        if font_dir.exists():
            preferred_result = self._find_preferred_local_font(font_dir)
            if preferred_result:
                return preferred_result

            for font_file in font_dir.glob("*.ttf"):
                result = self._try_register_local_font(font_file.stem, font_file)
                if result:
                    self.logger.debug(f"Using available font: {result}")
                    return result

        # Try built-in CID fonts with good Unicode support
        cid_fonts = ["HYSMyeongJoStd-Medium", "STSong-Light", "HeiseiMin-W3"]
        for cid_font in cid_fonts:
            result = self._try_register_cid_font(cid_font)
            if result:
                self.logger.debug(f"Using CID font: {result}")
                return result

        self.logger.warning("No suitable Unicode font found, using default Helvetica")
        return "Helvetica"

    def _create_pdf_styles(self, font_name):
        """Create PDF styles for the document."""
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

        styles = getSampleStyleSheet()

        base_settings = {
            "fontName": font_name,
            "encoding": "utf8",
            "splitLongWords": True,
            "spaceShrinkage": 0.05,
            "leading": 14,
        }

        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            **base_settings,
        )

        url_style = ParagraphStyle(
            "URLHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=24,
            **base_settings,
        )

        content_style = ParagraphStyle(
            "ContentStyle", parent=styles["Normal"], **base_settings
        )

        return title_style, url_style, content_style

    def _format_line_for_pdf(self, line, content_style):
        """Format and escape a single line for PDF inclusion."""
        if not line.strip():
            return Spacer(1, 7)

        try:
            import unicodedata

            normalized_text = unicodedata.normalize("NFC", line.strip())

            escaped_line = (
                normalized_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

            return Paragraph(escaped_line, content_style)
        except Exception as e:
            self.logger.warning(f"Error formatting text for PDF: {str(e)}")

            try:
                sanitized = ""
                for char in line.strip():
                    if ord(char) < 128:
                        sanitized += char
                    else:
                        try:
                            sanitized += f"&#{ord(char)};"
                        except Exception:
                            sanitized += "?"

                return Paragraph(sanitized, content_style)
            except Exception as e2:
                self.logger.error(f"Complete failure formatting text: {str(e2)}")
                return Paragraph("[Text conversion error]", content_style)

    def _add_content_to_pdf(self, story, scraped_data, url_style, content_style):
        """Add scraped content to the PDF story."""
        for url, content_text in scraped_data.items():

            story.append(Paragraph(f"URL: {url}", url_style))

            content_lines = content_text.split("\n")
            for line in content_lines:
                story.append(self._format_line_for_pdf(line, content_style))

            story.append(Spacer(1, 22))

        return story

    def _save_as_pdf(self, output_path, timestamp):
        """Save scraped content as PDF file."""
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        filename = f"{output_path}/scraped_content_{timestamp}.pdf"

        font_name = "Helvetica"
        try:
            registered_font = self._register_unicode_font()
            if registered_font:
                font_name = registered_font
                self.logger.debug(f"Using font {font_name} for PDF generation")
        except Exception as e:
            self.logger.warning(f"Font registration failed: {str(e)}")

        doc = SimpleDocTemplate(
            filename,
            pagesize=letter,
            encoding="utf-8",
            initialFontName=font_name,
            initialFontSize=11,
            allowSplitting=True,
        )

        title_style, url_style, content_style = self._create_pdf_styles(font_name)

        story = []

        story.append(Paragraph("Scraped Content", title_style))

        story.append(
            Paragraph(
                f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                content_style,
            )
        )
        story.append(Spacer(1, 14))

        story = self._add_content_to_pdf(
            story, self.scraped_data, url_style, content_style
        )

        doc.build(story)
        self.logger.debug(f"Successfully saved PDF file: {filename}")
        return filename

    def _save_as_json(self, output_path, timestamp):
        """Save scraped content as JSON file."""
        filename = f"{output_path}/scraped_content_{timestamp}.json"

        json_data = []
        for url, content in self.scraped_data.items():
            json_data.append({url: content})

        with open(filename, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

        self.logger.debug(f"Successfully saved JSON file: {filename}")
        return filename

    def get_scraped_content_as_string(self) -> str:
        """Get all scraped content as a single string."""
        return self._get_all_content()

    def _prepare_scraping_config(self, content_selectors=None, excluded_selectors=None):
        """
        Prepare a scraping configuration dictionary based on provided parameters.

        Args:
            content_selectors: Custom CSS selectors for content extraction
            excluded_selectors: Custom CSS selectors to exclude

        Returns:
            dict: Configuration dictionary for the scraper
        """
        scraping_config = {}

        if content_selectors:
            scraping_config["content_selectors"] = content_selectors

        if excluded_selectors:
            scraping_config["excluded_selectors"] = excluded_selectors

        return scraping_config if scraping_config else None

    async def scrape_website(
        self,
        start_url: str,
        max_depth: int = 2,
        allowed_domains: Optional[list] = None,
        scraping_config: Optional[dict] = None,
    ) -> dict:
        """
        Scrape a website starting from the given URL using a subprocess for event loop safety.

        Args:
            start_url: URL to start scraping from
            max_depth: Maximum depth to follow links
            allowed_domains: List of allowed domains to scrape
            scraping_config: Dictionary with content_selectors and excluded_selectors
        """
        try:
            return await self.scrape_website_subprocess(
                start_url=start_url,
                max_depth=max_depth,
                allowed_domains=allowed_domains,
                scraping_config=scraping_config,
            )
        except Exception as e:
            logging.error(f"Error during subprocess scraping: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "scraped_data": {},
                "failed_urls": [start_url],
                "total_pages": 0,
            }

    def save_to_file(
        self,
        scraped_data: dict,
        output_type: str = "text",
        output_path: str = None,
    ) -> str:
        """
        Save scraped content to a file.

        Args:
            scraped_data: Dictionary containing scraped content
            output_type: Type of output file ('txt', 'html', 'docx', 'pdf', 'json')
            output_path: Path where to save the file

        Returns:
            str: Path to the saved file
        """
        if not scraped_data:
            return ""

        if output_path is None:
            output_path = "./files"

        temp_spider = ScrapySpider()
        temp_spider.scraped_data = scraped_data

        return temp_spider._save_scraped_content_to_file(output_type, output_path)

    def get_content_as_string(self, scraped_data: dict) -> str:
        """
        Get all scraped content as a single string.

        Args:
            scraped_data: Dictionary containing scraped content

        Returns:
            str: All content combined as a single string
        """
        if not scraped_data:
            return ""

        all_content = ""
        for url, content in scraped_data.items():
            all_content += f"URL: {url}\n\n{content}\n\n{'-'*80}\n\n"
        return all_content

    async def scrape_websites(
        self,
        urls: list[str],
        max_depth: int = 1,
        output_type: str = "string",
        output_path: str = None,
        vector_db_index: str = None,
        allowed_domains: list[str] = None,
        content_selectors: list[str] = None,
        excluded_selectors: list[str] = None,
        embedding_model_config: Model = None,
    ) -> tuple[bool, str, list[str], list[str], str | list]:
        """
        Scrape multiple websites and return results in the format expected by the API.

        Args:
            urls: List of URLs to scrape
            max_depth: Maximum depth to follow links (default: 1)
            output_type: Type of output ('string', 'text', 'html', 'docx', 'pdf', 'json', 'vector_db')
            output_path: Path where to save files (optional)
            vector_db_index: Vector database index name (optional)
            allowed_domains: List of allowed domains to scrape (optional)
            content_selectors: Custom CSS selectors for content extraction (optional)
            excluded_selectors: Custom CSS selectors to exclude (optional)
            embedding_model_config: Model configuration object for the embedding model (optional)

        Returns:
            tuple: (success, message, scraped_urls, failed_urls, content)
        """
        try:
            scraping_config = self._prepare_scraping_config(
                content_selectors, excluded_selectors
            )

            scraped_results = await self._process_multiple_urls(
                urls, max_depth, allowed_domains, scraping_config
            )

            content = await self._generate_output(
                scraped_results["scraped_data"],
                output_type,
                output_path,
                vector_db_index,
                embedding_model_config,
            )

            message = self._generate_response_message(scraped_results)

            return (
                scraped_results["success"],
                message,
                scraped_results["scraped_urls"],
                scraped_results["failed_urls"],
                content,
            )

        except Exception as e:
            error_msg = f"Error during bulk scraping: {str(e)}"
            logging.error(error_msg)
            return False, error_msg, [], urls, ""

    async def _process_multiple_urls(
        self,
        urls: list[str],
        max_depth: int,
        allowed_domains: list[str],
        scraping_config: dict = None,
    ) -> dict:
        """Process multiple URLs and collect results."""
        all_scraped_data = {}
        all_failed_urls = []
        all_scraped_urls = []

        for url in urls:
            try:
                url_allowed_domains = allowed_domains or [urlparse(url).netloc]

                result = await self.scrape_website(
                    start_url=url,
                    max_depth=max_depth,
                    allowed_domains=url_allowed_domains,
                    scraping_config=scraping_config,
                )

                if result["success"]:
                    scraped = result["scraped_data"]
                    if isinstance(scraped, dict):
                        all_scraped_data.update(scraped)
                        all_scraped_urls.extend(scraped.keys())
                    elif isinstance(scraped, list):
                        for item in scraped:
                            if isinstance(item, dict):
                                all_scraped_data.update(item)
                                all_scraped_urls.extend(item.keys())
                    all_failed_urls.extend(result["failed_urls"])
                else:
                    all_failed_urls.append(url)
                    logging.error(
                        f"Failed to scrape {url}: {result.get('error', 'Unknown error')}"
                    )

            except Exception as e:
                all_failed_urls.append(url)
                logging.error(f"Error scraping {url}: {str(e)}")

        return {
            "success": len(all_scraped_data) > 0,
            "scraped_data": all_scraped_data,
            "scraped_urls": all_scraped_urls,
            "failed_urls": all_failed_urls,
        }

    async def _generate_output(
        self,
        scraped_data: dict,
        output_type: str,
        output_path: str,
        vector_db_index: str,
        embedding_model_config: Model = None,
    ) -> str | list:
        """Generate output content based on the specified type."""
        if output_type.lower() == OutputType.STRING.value:
            return self.get_content_as_string(scraped_data)
        elif output_type.lower() == OutputType.JSON.value:
            if output_path:
                file_path = self.save_to_file(
                    scraped_data, output_type.lower(), output_path
                )
                return f"Content saved to: {file_path}"
            else:
                json_data = []
                for url, content in scraped_data.items():
                    json_data.append({url: content})
                return json_data
        elif output_type.lower() == OutputType.VECTOR_DB.value:
            await self.save_to_vector_db(
                scraped_data, vector_db_index, embedding_model_config
            )
            return f"Content saved to vector database index: {vector_db_index or 'default'}"
        elif output_path and output_type.lower() in [
            OutputType.TXT.value,
            OutputType.HTML.value,
            OutputType.DOCX.value,
            OutputType.PDF.value,
        ]:
            file_path = self.save_to_file(
                scraped_data, output_type.lower(), output_path
            )
            return f"Content saved to: {file_path}"
        else:
            return self.get_content_as_string(scraped_data)

    def _generate_response_message(self, results: dict) -> str:
        """Generate a response message based on scraping results."""
        total_scraped = len(results["scraped_urls"])
        total_failed = len(results["failed_urls"])

        if results["success"]:
            message = f"Successfully scraped {total_scraped} page(s)"
            if total_failed > 0:
                message += f", {total_failed} failed"
            return message
        else:
            return f"Failed to scrape any content. {total_failed} URL(s) failed"

    def _extract_json_from_output(self, output: str):
        """
        Extract the first valid JSON object or array from a string that may contain extra log lines.
        """
        import json

        # Try line-by-line parsing first
        lines = output.splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                continue
            try:
                return json.loads(line)
            except Exception:
                continue
        # If not found, try bracket matching as fallback
        for start_char, end_char in [("{", "}"), ("[", "]")]:
            start = output.find(start_char)
            if start != -1:
                count = 0
                for i in range(start, len(output)):
                    if output[i] == start_char:
                        count += 1
                    elif output[i] == end_char:
                        count -= 1
                        if count == 0:
                            try:
                                return json.loads(output[start : i + 1])
                            except Exception:
                                break
        raise ValueError("No valid JSON found in output")

    async def scrape_website_subprocess(
        self,
        start_url: str,
        max_depth: int = 1,
        allowed_domains: Optional[list] = None,
        scraping_config: Optional[dict] = None,
    ) -> dict:
        """
        Run the Scrapy spider in a subprocess for robust event loop isolation.
        """
        script_path = str(Path(__file__).parent / "scrapy_subprocess_runner.py")
        allowed_domains_str = "" if not allowed_domains else ",".join(allowed_domains)

        cmd = [
            sys.executable,
            script_path,
            start_url,
            str(max_depth),
            allowed_domains_str,
        ]

        if scraping_config:

            selectors_json = json.dumps(scraping_config)
            selectors_b64 = base64.b64encode(selectors_json.encode("utf-8")).decode(
                "utf-8"
            )
            cmd.append(selectors_b64)

        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await proc.communicate()

        logging.info(f"Subprocess return code: {proc.returncode}")
        logging.info(f"Subprocess stdout: {stdout.decode()}")
        logging.info(f"Subprocess stderr: {stderr.decode()}")

        if proc.returncode != 0:
            return {
                "success": False,
                "error": stderr.decode(),
                "scraped_data": {},
                "failed_urls": [start_url],
                "total_pages": 0,
            }
        try:
            decoded_stdout = stdout.decode().strip()

            # First check if temp file exists at the specified location
            temp_file_path = str(
                Path(__file__).parent.parent.parent.parent / "files" / "temp_file.json"
            )

            if os.path.exists(temp_file_path):
                logging.info(f"Reading scraped data from temp file: {temp_file_path}")
                with open(temp_file_path, "r", encoding="utf-8") as temp_file:
                    scraped_data = json.load(temp_file)
                logging.info(
                    f"Successfully read {len(scraped_data)} items from temp file"
                )
            elif decoded_stdout.endswith(".json"):
                output_file_path = decoded_stdout
                logging.info(
                    f"Reading scraped data from output file: {output_file_path}"
                )
                with open(output_file_path, "r", encoding="utf-8") as output_file:
                    scraped_data = json.load(output_file)

                # Don't delete the file immediately - keep it for debugging
                logging.info(
                    f"Successfully read {len(scraped_data)} items from {output_file_path}"
                )
            else:
                scraped_data = self._extract_json_from_output(decoded_stdout)
        except Exception as e:
            return {
                "success": False,
                "error": f"JSON decode error: {e}",
                "scraped_data": {},
                "failed_urls": [start_url],
                "total_pages": 0,
            }
        return {
            "success": True,
            "scraped_data": scraped_data,
            "failed_urls": [],
            "total_pages": len(scraped_data),
            "start_url": start_url,
            "max_depth": max_depth,
            "allowed_domains": allowed_domains,
        }

    async def save_to_vector_db(
        self,
        scraped_data: dict,
        vector_db_index: str,
        embedding_model_config: Model = None,
    ) -> None:
        """
        Save scraped content to a Redis vector database.

        Args:
            scraped_data: Dictionary containing scraped content
            vector_db_index: Name of the Redis vector database index
            embedding_model_config: Model configuration object for the embedding model (optional)
        """
        if not scraped_data or not vector_db_index:
            return

        REDIS_URL = f"redis://{os.getenv('REDIS_USER', 'default')}:{os.getenv('REDIS_PASSWORD', '')}@{os.getenv('REDIS_HOST', '172.17.0.1')}:{os.getenv('REDIS_PORT', '6380')}"

        # Configure embedding model based on provided config or use defaults
        if embedding_model_config:
            provider = embedding_model_config.provider.value
            deployment = embedding_model_config.deployment
            model_name = embedding_model_config.name
            embedding_model = await get_embedding_model(
                provider=provider, deployment=deployment, model=model_name
            )
        else:
            embedding_model = await get_embedding_model()

        redis_config = RedisConfig(
            index_name=vector_db_index,
            redis_url=REDIS_URL,
            metadata_schema=[
                {"name": "timestamp", "type": "datetime"},
                {"name": "content", "type": "text"},
                {"name": "source", "type": "string"},
            ],
        )

        vector_store = RedisVectorStore(
            embedding_model, redis_config, index_name=vector_db_index
        )

        for url, content in scraped_data.items():
            await vector_store.add_texts(
                [content], metadatas=[{"source": url, "timestamp": datetime.now()}]
            )
